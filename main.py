import os
import logging
import re
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger()

def get_available_width(section):
    """Calculate available width in characters based on page settings."""
    # Get page width in inches
    page_width = section.page_width.inches
    # Subtract margins
    available_width = page_width - section.left_margin.inches - section.right_margin.inches + 2.7
    # Convert to approximate character count (assuming Courier New where each char is ~0.1 inches)
    char_count = int(available_width / 0.1)
    logger.debug(f"Calculated available width: {available_width} inches, approximately {char_count} characters.")
    return char_count

def split_text_into_lines(text, max_length):
    """Split text into lines with maximum length while preserving words."""
    lines = []
    words = text.split()
    current_line = []
    current_length = 0

    for word in words:
        if current_length + len(word) + (1 if current_length > 0 else 0) <= max_length:
            current_line.append(word)
            current_length += len(word) + (1 if current_length > 0 else 0)
        else:
            if current_line:
                lines.append(' '.join(current_line))
            current_line = [word]
            current_length = len(word)

    if current_line:
        lines.append(' '.join(current_line))

    logger.debug(f"Split content into {len(lines)} lines.")
    return lines

def add_row_content(table, text, is_separator=False, is_bold=False):
    """Add a full-width row for separators and segment markers."""
    row = table.add_row()
    if is_separator:
        cell = row.cells[0].merge(row.cells[1])
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.name = 'Verdana'
        run.font.size = Pt(11)
        if is_bold:
            run.font.bold = True
        paragraph.paragraph_format.space_after = Pt(6)
        logger.debug(f"Added separator line: {text}")
        return row
    logger.debug("Added row content.")
    return row

def add_split_content(table, speaker, content, speaker_color=None):
    """Add content split between speaker and content columns."""
    row = table.add_row()

    # Speaker cell
    if speaker:
        speaker_para = row.cells[0].paragraphs[0]
        speaker_run = speaker_para.add_run(speaker)
        if speaker_color:
            speaker_run.font.color.rgb = speaker_color
        speaker_para.paragraph_format.space_after = Pt(6)

    # Content cell
    if content:
        content_para = row.cells[1].paragraphs[0]
        content_para.add_run(content)
        content_para.paragraph_format.space_after = Pt(6)

    logger.debug(f"Added split content for speaker: {speaker}, content: {content[:30]}...")  # Log first 30 characters for brevity

def process_docx(input_file, output_file):
    if not os.path.exists(input_file):
        logger.error(f"Input file '{input_file}' not found.")
        raise FileNotFoundError(f"Input file '{input_file}' not found.")

    logger.info(f"Processing document: {input_file}")
    doc = Document(input_file)
    out_doc = Document()

    # Set document defaults
    style = out_doc.styles['Normal']
    style.font.name = 'Verdana'
    style.font.size = Pt(12)

    # Set page margins
    section = out_doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Calculate available width in characters
    available_width = get_available_width(section)
    separator_length = available_width + 4  # Leave some margin
    logger.debug(f"Separator length set to: {separator_length}")

    # Create main table
    table = out_doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.allow_autofit = False

    # Set column widths
    table.columns[0].width = Cm(3)  # Speaker column
    table.columns[1].width = Cm(16)  # Content column

    # Make table borders invisible
    for row in table.rows:
        for cell in row.cells:
            cell._tc.tcPr.tcBorders.top.val = 'nil'
            cell._tc.tcPr.tcBorders.bottom.val = 'nil'
            cell._tc.tcPr.tcBorders.left.val = 'nil'
            cell._tc.tcPr.tcBorders.right.val = 'nil'

    # Define constants
    MAX_LINE_LENGTH = 90  # Maximum characters per line for content

    speaker_pattern = re.compile(r'^[A-Z]+:?\s')
    segment_marker_pattern = re.compile(r'\b(\d{1,2}:\d{2}(?::\d{2})?)\b')
    timecode_pattern = re.compile(r'\b(\d{1,2}:\d{2}(?::\d{2})?)\b')

    segment_count = 0
    skip_pages = True

    for para in doc.paragraphs:
        text = para.text.strip()

        if skip_pages:
            if segment_marker_pattern.search(text) and "----------" in text:
                skip_pages = False
            else:
                continue

        # Handle segment markers
        if segment_marker_pattern.search(text) and "----------" in text:
            segment_count += 1
            # Add separator line
            add_row_content(table, '-' * separator_length, is_separator=True)

            # Extract timecode without dashes
            timecode = segment_marker_pattern.search(text).group(1)
            # Calculate spacing dynamically
            spacing = separator_length - len(timecode) - 2 + 26  # 2 for segment number width
            # Adjusted spacing calculation
            segment_text = f"{timecode}" + " " * spacing + f" {segment_count:02d}"

            add_row_content(table, segment_text, is_separator=True, is_bold=True)
            logger.info(f"Added segment marker: {segment_text}")

            continue

        # Remove timecode if present
        if timecode_pattern.match(text):
            timecode = timecode_pattern.match(text).group(0)
            text = text.replace(timecode, "").strip()

        # Process speakers and content
        if speaker_pattern.match(text):
            words = text.split()
            speaker = words[0].rstrip(':') if words else ""  # Remove colon if present

            content = ' '.join(words[1:]).strip()

            if content:
                # Split content into lines
                lines = split_text_into_lines(content, MAX_LINE_LENGTH)

                # Add first line with speaker
                add_split_content(table, speaker, lines[0], speaker_color=RGBColor(255, 0, 0))

                # Add remaining lines
                for line in lines[1:]:
                    add_split_content(table, "", line)
            else:
                # Handle speaker-only lines
                add_split_content(table, speaker, "", speaker_color=RGBColor(255, 0, 0))
            logger.info(f"Processed speaker: {speaker}, content: {content[:30]}...")  # Log first 30 characters for brevity
        else:
            # Handle non-speaker lines (like TITULOK)
            if text:
                add_split_content(table, "", text)
            logger.debug(f"Processed non-speaker line: {text[:30]}...")

    # Ensure the output file is saved as .docx
    output_file = output_file if output_file.endswith('.docx') else output_file + '.docx'

    # Save the processed document
    out_doc.save(output_file)
    logger.info(f"Processed document saved as {output_file}")

    return output_file

if __name__ == "__main__":
    input_file = sys.argv[1]
    output_file = "output.docx"  # Example output path

    result = process_docx(input_file, output_file)
    if result:
        print(result)  # Return the processed file path as output
    else:
        print("Error: Could not process the file.")
        